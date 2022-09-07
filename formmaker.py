import os
import docx
from docx.shared import Cm, Inches
from datetime import datetime

# paragraph_format.left_indent = Inches(0.5)
# paragraph_format.first_line_indent
#Set font
# font = run.font
# font.size = Pt(16)

DATE = datetime.today().strftime("%m/%d/%Y")

#change directory into save folder /docs
os.chdir(os.path.join(os.getcwd(), "docs"))

document_number = 1

not_finished = True
while(not_finished):

    # Create instance of word document
    doc = docx.Document()

    # Write the first section of the document
    doc_header = doc.add_paragraph()
    header_run = doc_header.add_run("EQUIPMENT CHAIN OF CUSTODY")
    header_run.bold = True
    header_run.add_break()


    ###
    sect1 = doc.add_paragraph("")
    sect1_run = sect1.add_run("Date:" + DATE)
    sect1_run.bold = True
    sect1_run = sect1.add_run("I ________________________________ ACCEPT CUSTODY OF THE PROPERTY BEING")
    sect1_run = sect1.add_run("DELIVERED BY ________________________________________  LISTED BELOW:")


    # Create the table 
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    # table.allow_autofit = False
    table.columns[0].width = Inches(1.0)
    table.rows[0].cells[0].width = Inches(1.0)


    #Fill in First Row
    i = 0
    for cell in table.rows[0].cells: 
        paragraph = cell.paragraphs[0]
        if i == 0:
            run = paragraph.add_run("Bitlocker")
        elif i == 1:
            run = paragraph.add_run("Serial Number")
        elif i == 2:
            run = paragraph.add_run("Employee Name / Job Title")
        elif i == 3:    
            run = paragraph.add_run("Description")
        run.bold = True     
        i = i + 1

    #for i in range(4):
        #for cell in table.rows[i].cells:
            #if i == 2:
               # cell.width = Inches(1)
               # cell.height = Inches (0.25)
            #else:
                #cell.width = Inches(0.5)

    # Get the inputs for the table
    for i in range(4):
        bitlocker = input("Enter the bitlocker: ")
        serial_number = input("Enter the serial number: ")
        employee_name = input("Enter the employee name: ")
        job_title = input("Enter the job_title: ")
        description = input("Enter the description: ")

        print("-----------------------------------")

        name_desc = employee_name + "/" + job_title
        row = table.rows[i+1].cells
        row[0].text = bitlocker;
        row[1].text = serial_number;
        row[2].text = name_desc;
        row[3].text = description;



    # Write the third section of the document
     # Write the third section of the document
    doc_paragraph2 = doc.add_paragraph("")
    par2 = doc_paragraph2.add_run()
    par2.add_break()
    doc.add_paragraph("I UNDERSTAND ______________________________________ IS ACCOUNTABLE FOR THIS ")
    doc.add_paragraph("A erat nam at lectus urna duis. Tortor condimentum lacinia quis vel eros donec ac odio. Maecenas pharetra convallis posuere morbi. Interdum posuere lorem ipsum dolor sit amet consectetur adipiscing.")
    doc.add_paragraph("Sign here ________________________________________________________")
    doc.add_paragraph("Sign here ________________________________________________________")


    # Save
    save_name = "\\legal_hold" + "_" + DATE + "(" + str(document_number) + ")" + ".docx"
    doc.save(save_name)

    # Check if the user would like to continue
    not_finished = input("Would you like to continue (y/n): ")
    not_finished = not_finished.lower()
    if (not_finished == "yes" or not_finished == "y"):
        document_number += 1;
        pass
    elif (not_finished == "no" or not_finished == "n"):
        not_finished = False
    else:
        not_finished = False

